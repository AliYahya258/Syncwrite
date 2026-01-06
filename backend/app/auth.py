from fastapi import APIRouter, HTTPException, Depends, Header, UploadFile, File
from fastapi.responses import StreamingResponse
from typing import Optional
from .models import (
    RegisterRequest, LoginRequest, LoginResponse,
    CreateRoomRequest, RoomResponse,
    InviteUserRequest, InvitationResponse, AcceptInviteRequest
)
from .db import (
    create_user, verify_user, get_user_by_id, get_user_by_email,
    create_room, get_user_rooms, check_room_access, get_room_users,
    create_invitation, get_user_invitations, accept_invitation, decline_invitation,
    grant_room_access, revoke_room_access,
    get_all_users, get_all_rooms, delete_user_admin, delete_room_admin,
    make_user_admin, check_is_admin, update_document_content
)
from .jwt_utils import create_access_token, verify_token
from .redis_client import r
import io
from pydantic import BaseModel

router = APIRouter()


# Helper function to get current user from token
async def get_current_user(authorization: Optional[str] = Header(None)) -> dict:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")
    
    token = authorization.split(" ")[1]
    user_data = verify_token(token)
    
    if not user_data:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    return user_data


@router.post("/api/register")
async def register(req: RegisterRequest):
    try:
        user_id = create_user(req.username, req.password, req.email)
        token = create_access_token(user_id, req.email, req.username)
        return {
            "user_id": user_id,
            "username": req.username,
            "email": req.email,
            "token": token,
            "is_admin": False
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/api/login")
async def login(req: LoginRequest) -> LoginResponse:
    result = verify_user(req.email, req.password)
    if not result:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    user_id, username, email, is_admin = result
    token = create_access_token(user_id, email, username)
    
    return LoginResponse(
        user_id=user_id,
        username=username,
        email=email,
        token=token,
        is_admin=is_admin
    )


@router.get("/api/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    """Get current user info"""
    # Add admin status from database
    is_admin = check_is_admin(current_user["user_id"])
    return {
        **current_user,
        "is_admin": is_admin
    }


@router.post("/api/rooms")
async def create_new_room(
    req: CreateRoomRequest,
    current_user: dict = Depends(get_current_user)
) -> RoomResponse:
    """Create a new room"""
    try:
        room_id = create_room(current_user["user_id"], req.room_name)
        return RoomResponse(
            room_id=room_id,
            room_name=req.room_name,
            owner_id=current_user["user_id"],
            role="owner"
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/api/rooms")
async def list_rooms(current_user: dict = Depends(get_current_user)):
    """List all rooms user has access to"""
    rooms = get_user_rooms(current_user["user_id"])
    return {"rooms": rooms}


@router.delete("/api/rooms/{room_id:path}")
async def delete_own_room(
    room_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a room (owner only)"""
    # Check if current user is the owner
    role = check_room_access(room_id, current_user["user_id"])
    if role != "owner":
        raise HTTPException(status_code=403, detail="Only the owner can delete a room")
    
    # Use the admin delete function since it does the same thing
    delete_room_admin(room_id)
    return {"message": "Room deleted successfully"}


@router.get("/api/rooms/{room_id:path}/users")
async def get_room_users_endpoint(
    room_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all users with access to a room"""
    # Check if user has access to the room
    role = check_room_access(room_id, current_user["user_id"])
    if not role:
        raise HTTPException(status_code=403, detail="Access denied")
    
    users = get_room_users(room_id)
    
    # Get active users from Redis
    active_user_ids = list(r.smembers(f"presence:{room_id}"))
    
    # Mark active users
    for user in users:
        user["is_online"] = user["user_id"] in active_user_ids
    
    return {"room_id": room_id, "users": users, "count": len(users)}


@router.post("/api/invitations")
async def invite_user(
    req: InviteUserRequest,
    current_user: dict = Depends(get_current_user)
):
    """Invite a user to a room"""
    # Check if current user is owner or editor
    role = check_room_access(req.room_id, current_user["user_id"])
    if role not in ["owner", "editor"]:
        raise HTTPException(status_code=403, detail="Only owners and editors can invite users")
    
    # Check if invited user exists
    invited_user = get_user_by_email(req.invited_email)
    if not invited_user:
        raise HTTPException(status_code=404, detail="User with this email not found")
    
    # Check if user already has access
    existing_access = check_room_access(req.room_id, invited_user["user_id"])
    if existing_access:
        raise HTTPException(status_code=400, detail="User already has access to this room")
    
    invite_id = create_invitation(
        req.room_id,
        current_user["user_id"],
        req.invited_email,
        req.role
    )
    
    return {"invite_id": invite_id, "message": "Invitation sent"}


@router.get("/api/invitations")
async def list_invitations(current_user: dict = Depends(get_current_user)):
    """Get all pending invitations for current user"""
    invitations = get_user_invitations(current_user["email"])
    return {"invitations": invitations}


@router.post("/api/invitations/{invite_id}/accept")
async def accept_invite(
    invite_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Accept an invitation"""
    result = accept_invitation(invite_id, current_user["user_id"])
    if not result:
        raise HTTPException(status_code=404, detail="Invitation not found or already processed")
    
    return {"message": "Invitation accepted", "room_id": result["room_id"], "role": result["role"]}


@router.post("/api/invitations/{invite_id}/decline")
async def decline_invite(
    invite_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Decline an invitation"""
    decline_invitation(invite_id)
    return {"message": "Invitation declined"}


@router.delete("/api/rooms/{room_id:path}/users/{user_id}")
async def remove_user_from_room(
    room_id: str,
    user_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Remove a user from a room (owner only)"""
    # Check if current user is owner
    role = check_room_access(room_id, current_user["user_id"])
    if role != "owner":
        raise HTTPException(status_code=403, detail="Only the owner can remove users")
    
    # Cannot remove the owner
    target_role = check_room_access(room_id, user_id)
    if target_role == "owner":
        raise HTTPException(status_code=400, detail="Cannot remove the owner")
    
    revoke_room_access(room_id, user_id)
    return {"message": "User removed from room"}


@router.put("/api/rooms/{room_id:path}/users/{user_id}/role")
async def update_user_role(
    room_id: str,
    user_id: str,
    role: str,
    current_user: dict = Depends(get_current_user)
):
    """Update a user's role in a room (owner only)"""
    if role not in ["editor", "viewer"]:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    # Check if current user is owner
    current_role = check_room_access(room_id, current_user["user_id"])
    if current_role != "owner":
        raise HTTPException(status_code=403, detail="Only the owner can change roles")
    
    # Cannot change the owner's role
    target_role = check_room_access(room_id, user_id)
    if target_role == "owner":
        raise HTTPException(status_code=400, detail="Cannot change the owner's role")
    
    grant_room_access(room_id, user_id, role, current_user["user_id"])
    return {"message": f"User role updated to {role}"}


# ============ Admin Endpoints ============

async def get_admin_user(current_user: dict = Depends(get_current_user)) -> dict:
    """Verify that the current user is an admin"""
    is_admin = check_is_admin(current_user["user_id"])
    print(f"Admin check for user {current_user['user_id'][:8]}...: is_admin={is_admin}")
    if not is_admin:
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user


@router.get("/api/admin/users")
async def admin_get_all_users(admin_user: dict = Depends(get_admin_user)):
    """Get all users (admin only)"""
    users = get_all_users()
    return {"users": users}


@router.get("/api/admin/rooms")
async def admin_get_all_rooms(admin_user: dict = Depends(get_admin_user)):
    """Get all rooms (admin only)"""
    rooms = get_all_rooms()
    return {"rooms": rooms}


@router.delete("/api/admin/users/{user_id}")
async def admin_delete_user(
    user_id: str,
    admin_user: dict = Depends(get_admin_user)
):
    """Delete a user and all their rooms (admin only)"""
    if user_id == admin_user["user_id"]:
        raise HTTPException(status_code=400, detail="Cannot delete yourself")
    
    delete_user_admin(user_id)
    return {"message": "User deleted successfully"}


@router.delete("/api/admin/rooms/{room_id:path}")
async def admin_delete_room(
    room_id: str,
    admin_user: dict = Depends(get_admin_user)
):
    """Delete a room (admin only)"""
    delete_room_admin(room_id)
    return {"message": "Room deleted successfully"}


@router.post("/api/admin/users/{user_id}/make-admin")
async def admin_make_admin(
    user_id: str,
    admin_user: dict = Depends(get_admin_user)
):
    """Promote a user to admin (admin only)"""
    make_user_admin(user_id)
    return {"message": "User promoted to admin successfully"}


# ============ File Upload/Download Endpoints ============

class DownloadRequest(BaseModel):
    room_id: str
    content: str


@router.post("/api/upload-docx")
async def upload_docx(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a DOCX file and convert to HTML"""
    try:
        from docx import Document
        import mammoth
        
        # Read the uploaded file
        contents = await file.read()
        
        # Convert DOCX to HTML using mammoth
        result = mammoth.convert_to_html(io.BytesIO(contents))
        html_content = result.value
        
        return {"html_content": html_content, "messages": result.messages}
    except ImportError:
        raise HTTPException(
            status_code=500, 
            detail="Document conversion libraries not installed. Run: pip install python-docx mammoth"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@router.post("/api/download-pdf")
async def download_pdf(
    req: DownloadRequest,
    current_user: dict = Depends(get_current_user)
):
    """Download document as PDF"""
    try:
        from weasyprint import HTML
        
        # Check if user has access to the room
        role = check_room_access(req.room_id, current_user["user_id"])
        if not role:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Create PDF from HTML content
        html_with_styles = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    padding: 2cm;
                    max-width: 21cm;
                    margin: 0 auto;
                }}
                h1, h2, h3 {{ margin-top: 1em; margin-bottom: 0.5em; }}
                p {{ margin: 0.5em 0; }}
            </style>
        </head>
        <body>
            {req.content}
        </body>
        </html>
        """
        
        pdf_bytes = HTML(string=html_with_styles).write_pdf()
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=document.pdf"}
        )
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF generation library not installed. Run: pip install weasyprint"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")


@router.post("/api/download-docx")
async def download_docx(
    req: DownloadRequest,
    current_user: dict = Depends(get_current_user)
):
    """Download document as DOCX"""
    try:
        from docx import Document
        from bs4 import BeautifulSoup
        
        # Check if user has access to the room
        role = check_room_access(req.room_id, current_user["user_id"])
        if not role:
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Create a new Document
        doc = Document()
        
        # Parse HTML content
        soup = BeautifulSoup(req.content, 'html.parser')
        
        # Convert HTML to DOCX paragraphs
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'li']):
            text = element.get_text()
            if element.name == 'h1':
                doc.add_heading(text, level=1)
            elif element.name == 'h2':
                doc.add_heading(text, level=2)
            elif element.name == 'h3':
                doc.add_heading(text, level=3)
            elif element.name == 'p':
                doc.add_paragraph(text)
            elif element.name in ['li']:
                doc.add_paragraph(text, style='List Bullet')
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=document.docx"}
        )
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Document libraries not installed. Run: pip install python-docx beautifulsoup4"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


